import asyncio, time
from dataclasses import dataclass, field
from typing import Optional
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
import uvicorn
import os, re, uuid, tempfile, json, shutil
import PyPDF2
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from config import MAX_FILE_SIZE
import fitz
from services.llm_service import (
    analyze_document, generate_question, generate_paper_stream, edit_paper,
    analyze_document_for_resume, generate_resume_question, generate_resume_stream, edit_resume
)

# ─── Paper Format presets ──────────────────────────────────────────────────────

@dataclass
class PaperFormat:
    name: str
    page_size_mm: tuple = (210, 297)
    margins_mm: tuple = (18, 18, 18, 18)
    column_gap_mm: float = 6
    header_text: Optional[str] = None
    header_font_size: int = 9
    footer_text: Optional[str] = None
    footer_font_size: int = 9
    title_font_size: int = 20
    author_font_size: int = 12
    body_font_size: int = 10
    section_font_size: int = 10
    abstract_bold_italic: bool = True
    section_centered: bool = True
    font_family: str = "Times"

PROCOMM = PaperFormat(
    name="IEEE ProComm",
    abstract_bold_italic=True,
    section_centered=True,
)

IEMT = PaperFormat(
    name="IEMT",
    footer_text="36th International Electronic Manufacturing Technology Conference, 2014",
    footer_font_size=10,
    abstract_bold_italic=False,
    section_centered=False,
    title_font_size=12,
    author_font_size=10,
)

app = FastAPI(title="Research Paper Generator")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

SESSION_TTL_MINUTES = 60
sessions = {}
_last_session_access = {}

def _touch_session(sid: str):
    _last_session_access[sid] = time.time()

def _session_cleanup():
    while True:
        now = time.time()
        stale = [sid for sid, t in _last_session_access.items() if now - t > SESSION_TTL_MINUTES * 60]
        for sid in stale:
            s = sessions.pop(sid, None)
            _last_session_access.pop(sid, None)
        time.sleep(300)

@app.on_event("startup")
async def _start_cleaner():
    import threading
    t = threading.Thread(target=_session_cleanup, daemon=True)
    t.start()

def _render_body(content, session_id=""):
    """Turn raw section text into HTML: paragraphs, markdown tables, captions, images."""
    def esc(t):
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    def is_sep(line):
        s = line.strip().strip("|")
        return bool(s) and set(s) <= set("-: |") and "-" in s
    def is_caption(s):
        return bool(re.match(r'^(figure|table)\s*\d*\.?\s', s, re.I))
    def md_table(rows):
        def cells(r):
            return [c.strip() for c in r.strip().strip("|").split("|")]
        header = cells(rows[0])
        body = [cells(r) for r in rows[2:] if "|" in r]
        th = "".join(f'<th class="tablehead">{esc(c)}</th>' for c in header)
        trs = "".join(
            "<tr>" + "".join(f'<td class="tabletext">{esc(c)}</td>' for c in r) + "</tr>"
            for r in body
        )
        return f'<table class="ieee-table"><thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table>'
    img_re = re.compile(r'!\[(.*?)\]\(([^)]+)\)')

    lines = [l.rstrip() for l in content.split("\n")]
    out, i, n = [], 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        img_m = img_re.match(stripped)
        if img_m:
            alt = esc(img_m.group(1))
            fname = img_m.group(2)
            src = f"/api/session/{session_id}/image/{fname}" if session_id else fname
            out.append(f'<div class="figure"><img src="{src}" alt="{alt}" />')
            out.append(f'<div class="figcaption">{alt}</div></div>')
            i += 1
            continue
        if "|" in stripped and i + 1 < n and is_sep(lines[i + 1]):
            rows = []
            while i < n and "|" in lines[i].strip():
                rows.append(lines[i].strip())
                i += 1
            out.append(md_table(rows))
            continue
        if is_caption(stripped):
            out.append(f'<div class="caption">{esc(stripped)}</div>')
            i += 1
            continue
        para = [stripped]
        i += 1
        while i < n and lines[i].strip() and "|" not in lines[i] and not is_sep(lines[i]) and not is_caption(lines[i].strip()):
            img_m2 = img_re.match(lines[i].strip())
            if img_m2:
                break
            para.append(lines[i].strip())
            i += 1
        out.append(f'<p>{esc(" ".join(para))}</p>')
    return "".join(out)

def generate_ieee_html(title, authors, abstract, sections, keywords, domain, references=None, session_id=""):
    # Build author HTML with per-affiliation superscript markers
    authors_html = ""
    if authors:
        affil_map = {}
        counter = 1
        author_parts = []
        for a in authors:
            affil = a.get("affiliation", "")
            if affil:
                if affil not in affil_map:
                    affil_map[affil] = counter
                    counter += 1
                sup = f'<sup>{affil_map[affil]}</sup>'
            else:
                sup = ""
            email = a.get("email", "")
            email_html = f'<br><span class="email">{email}</span>' if email else ""
            author_parts.append(
                f'<div class="author">{a["name"]}{sup}<br><span class="affil">{affil}</span>{email_html}</div>'
            )
        authors_html = "".join(author_parts)

    sections_html = "".join(
        f'<div class="section"><h2>{s["title"]}</h2>{_render_body(s["content"], session_id)}</div>'
        for s in sections
    )
    keywords_str = ", ".join(keywords)

    # Build references HTML
    refs_html = ""
    if references:
        ref_list = "\n".join(
            f'<p>[{i}] {r["citation"]}</p>' for i, r in enumerate(references, 1)
        )
        refs_html = f'<div class="references"><h2>References</h2>\n{ref_list}\n</div>'

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>{title}</title>
<style>
  @page {{ size: letter; margin: 0.75in 0.65in; }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: "Times New Roman", Times, serif; font-size: 10pt; line-height: 1.15; color: #000; background: #fff; }}
  .paper {{ max-width: 8.5in; margin: 0 auto; padding: 0.75in 0.65in; }}
  h1 {{ font-size: 24pt; text-align: center; font-weight: bold; margin-bottom: 16px; }}
  .authors {{ text-align: center; font-size: 12pt; margin-bottom: 18px; }}
  .author {{ display: inline-block; margin: 0 16px; }}
  .affil {{ font-size: 10pt; font-style: italic; }}
  .email {{ font-size: 10pt; }}
  sup {{ font-size: 8pt; vertical-align: super; line-height: 1; }}
  .abstract {{ margin: 12px 0; padding: 0; }}
  .abstract-label {{ font-size: 10pt; font-weight: bold; font-style: italic; }}
  .abstract p {{ font-size: 10pt; font-weight: bold; font-style: italic; text-align: justify; display: inline; }}
  .kw-label {{ font-size: 10pt; font-weight: bold; font-style: italic; }}
  .keywords {{ font-size: 10pt; margin-bottom: 12px; font-style: italic; }}
  .header {{ column-span: all; width: 100%; page-break-after: avoid; }}
  .content {{ column-count: 2; column-gap: 0.3in; column-rule: 0.5pt solid #ccc; }}
  .content::after {{ content: ""; display: table; clear: both; }}
  .section {{ margin-bottom: 0; break-inside: avoid-column; }}
  .section h2 {{ font-size: 10pt; font-variant: small-caps; font-weight: bold; text-align: center; margin: 12pt 0 6pt 0; letter-spacing: 0.5pt; }}
  .section h3 {{ font-size: 10pt; font-style: italic; font-weight: normal; text-align: left; margin: 9pt 0 3pt 0; }}
  .section p {{ text-align: justify; text-indent: 0.17in; margin-bottom: 0; line-height: 1.15; }}
  .figure {{ text-align: center; margin: 10pt 0; break-inside: avoid; }}
  .figure img {{ max-width: 100%; height: auto; }}
  .figcaption {{ font-size: 9pt; font-style: italic; text-align: center; margin-top: 4pt; }}
  .caption {{ font-size: 10pt; font-variant: small-caps; text-align: left; margin: 6pt 0; }}
  table.ieee-table {{ border-collapse: collapse; width: 100%; font-size: 9pt; margin: 8px 0; }}
  table.ieee-table th, table.ieee-table td {{ border: 0.5pt solid black; padding: 3px 6px; text-align: center; }}
  .tablehead {{ font-variant: small-caps; font-weight: bold; }}
  .tabletext {{ text-align: left; }}
  .references {{ margin-top: 12px; }}
  .references h2 {{ font-size: 10pt; font-variant: small-caps; font-weight: bold; text-align: center; margin: 12pt 0 6pt 0; }}
  .references p {{ font-size: 9pt; margin-left: 0.25in; text-indent: -0.25in; line-height: 1.15; margin-bottom: 6pt; text-align: left; }}
  @media screen and (max-width: 700px) {{ .content {{ column-count: 1; }} }}
  @media print {{ body {{ padding: 0; }} .paper {{ padding: 0; }} }}
</style></head><body>
<div class="paper">
  <div class="header">
    <h1>{title}</h1>
    <div class="authors">{authors_html}</div>
    <div class="abstract"><span class="abstract-label">Abstract - </span><p>{abstract}</p></div>
    <div class="keywords"><span class="kw-label">Index Terms - </span>{keywords_str}</div>
  </div>
  <div class="content">{sections_html}
  {refs_html}</div>
</div>
</body></html>"""

def extract_text(file_path: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        text = ""
        with open(file_path, 'rb') as f:
            for page in PyPDF2.PdfReader(f).pages:
                text += page.extract_text() + "\n"
        return text
    elif ext == '.docx':
        text = ""
        for para in Document(file_path).paragraphs:
            text += para.text + "\n"
        return text
    elif ext == '.ipynb':
        with open(file_path, 'r', encoding='utf-8') as f:
            nb = json.load(f)
        text = ""
        for cell in nb.get('cells', []):
            src = ''.join(cell.get('source', []))
            if cell.get('cell_type') == 'code':
                text += f"[code]\n{src}\n[/code]\n"
            else:
                text += src + "\n"
        return text
    elif ext == '.html' or ext == '.htm':
        from html.parser import HTMLParser
        class TagStripper(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text = []
            def handle_data(self, data):
                self.text.append(data)
        stripper = TagStripper()
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            stripper.feed(f.read())
        return ''.join(stripper.text)
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


def extract_text_per_page(file_path: str) -> dict:
    """Extract text from a PDF keyed by page number (1-indexed)."""
    pages = {}
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages, 1):
                t = page.extract_text()
                if t.strip():
                    pages[i] = t.strip()
    except Exception:
        pass
    return pages


def extract_pdf_images(file_path: str, session_dir: str, page_texts: Optional[dict] = None) -> list:
    """Extract images from a PDF, save to session_dir/images/, return metadata with page context."""
    images_dir = os.path.join(session_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    images = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page_text = (page_texts or {}).get(page_num + 1, "")
            for img_index, img in enumerate(doc.get_page_images(page_num)):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.width < 100 or pix.height < 100:
                    continue
                if pix.n - pix.alpha > 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                ext = "png" if not img[2] else "jpeg"
                fname = f"img_p{page_num+1}_{img_index+1}.{ext}"
                fpath = os.path.join(images_dir, fname)
                pix.save(fpath)
                size_kb = os.path.getsize(fpath) / 1024
                if size_kb < 5:
                    os.remove(fpath)
                    continue
                context = page_text[:500] if page_text else ""
                images.append({
                    "filename": fname, "path": fpath,
                    "page": page_num + 1,
                    "width": pix.width, "height": pix.height,
                    "size_kb": round(size_kb, 1),
                    "context": context,
                })
                pix = None
        doc.close()
    except Exception:
        pass
    return images


def extract_docx_images(file_path: str, session_dir: str) -> list:
    """Extract images from a DOCX file."""
    images_dir = os.path.join(session_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    images = []
    try:
        doc = Document(file_path)
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.reltype:
                img = rel.target_part
                ext = img.content_type.split("/")[-1]
                if ext == "jpeg": ext = "jpg"
                fname = f"img_docx_{i+1}.{ext}"
                fpath = os.path.join(images_dir, fname)
                with open(fpath, "wb") as f:
                    f.write(img.blob)
                size_kb = os.path.getsize(fpath) / 1024
                if size_kb < 5:
                    os.remove(fpath)
                    continue
                images.append({
                    "filename": fname, "path": fpath,
                    "page": 0, "width": 0, "height": 0,
                    "size_kb": round(size_kb, 1),
                })
    except Exception:
        pass
    return images


@app.get("/")
async def root():
    return {"message": "Research Paper Generator", "status": "running"}

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(400, "File too large")

    tmp = tempfile.mkdtemp()
    safe_name = os.path.basename(file.filename or "upload")
    file_path = os.path.join(tmp, safe_name)
    with open(file_path, 'wb') as f:
        f.write(content)

    text = extract_text(file_path, file.filename)
    if not text.strip():
        raise HTTPException(400, "Could not extract text from file")

    ext = os.path.splitext(file.filename)[1].lower()
    images = []
    if ext == ".pdf":
        page_texts = extract_text_per_page(file_path)
        images = extract_pdf_images(file_path, tmp, page_texts)
    elif ext == ".docx":
        images = extract_docx_images(file_path, tmp)

    session_id = str(uuid.uuid4())
    sessions[session_id] = {
        "session_dir": tmp,
        "file_text": text, "filename": file.filename,
        "answers": {}, "questions_asked": [],
        "ready": False, "analysis": {},
        "paper_text": None,
        "images": images,
    }

    return {"session_id": session_id, "images_count": len(images)}

@app.get("/api/session/{session_id}/image/{filename}")
async def serve_session_image(session_id: str, filename: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)
    safe_name = os.path.basename(filename)
    if not safe_name:
        raise HTTPException(400, "Invalid filename")
    full_path = os.path.join(s["session_dir"], "images", safe_name)
    if not os.path.isfile(full_path):
        raise HTTPException(404, "Image not found")
    ext = os.path.splitext(safe_name)[1].lower()
    media_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".gif": "image/gif"}
    return Response(content=open(full_path, "rb").read(), media_type=media_map.get(ext, "image/png"))


@app.post("/api/ask-paper/{session_id}")
async def ask_paper_question(session_id: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    q_result = generate_question(s["file_text"], s["answers"], s["questions_asked"], s.get("analysis"))
    if q_result.get("ready"):
        s["ready"] = True
        return {"ready": True}

    s["_last_qtype"] = q_result.get("type", "")
    return {"question": q_result.get("question", ""), "options": q_result.get("options", []), 
            "context": q_result.get("context", ""), "type": q_result.get("type", "")}

@app.post("/api/answer-paper/{session_id}")
async def submit_paper_answer(session_id: str, data: dict):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    question = data.get("question", "")
    answer = data.get("answer", "")
    if not isinstance(question, str):
        return {"error": "question must be a string", "question": "", "options": [], "context": "", "type": ""}
    s["answers"][question] = answer
    s["questions_asked"].append(question)

    qtype = s.get("_last_qtype", "")

    if qtype in ("authors", "authors_confirm_no"):
        names = [n.strip() for n in answer.replace(",", ";").split(";") if n.strip()]
        if names:
            s["analysis"]["authors"] = names
        s["answers"]["_authors_ok"] = True
    elif qtype == "authors_confirm":
        if answer.lower().startswith("y"):
            s["answers"]["_authors_ok"] = True
        else:
            s["_last_qtype"] = "authors_confirm_no"
            return {
                "question": "Please type the correct author name(s), separated by semicolons.",
                "options": [],
                "context": "Correct names for the paper byline.",
                "type": "authors_confirm_no"
            }
    elif qtype == "affiliation":
        if answer.strip():
            s["analysis"]["affiliation"] = answer.strip()
        s["answers"]["_affiliation_ok"] = True
    elif qtype == "email":
        if answer.strip():
            s["analysis"]["contact_email"] = answer.strip()
        s["answers"]["_email_ok"] = True
    elif qtype == "title":
        if "I'll type" in answer or "type the title" in answer:
            s["answers"]["_need_typed_title"] = True
        elif "first line" in answer:
            first_line = (s["file_text"] or "").strip().split("\n")[0][:100]
            s["analysis"]["title"] = first_line
            s["answers"]["_title_ok"] = True
        elif answer.strip():
            s["analysis"]["title"] = answer.strip()
            s["answers"]["_title_ok"] = True
    elif qtype == "keywords":
        if answer.strip():
            kws = [k.strip() for k in answer.replace(";", ",").split(",") if k.strip()]
            s["analysis"]["keywords"] = kws
        s["answers"]["_keywords_ok"] = True
    elif qtype in ("title_confirm", "title_correct"):
        if qtype == "title_confirm" and not answer.lower().startswith("y"):
            s["_last_qtype"] = "title_correct"
            return {
                "question": "Please type the correct paper title.",
                "options": [s["analysis"].get("title", "")],
                "context": "The title at the top of the paper.",
                "type": "title_correct"
            }
        if qtype == "title_correct" and answer.strip():
            s["analysis"]["title"] = answer.strip()
        s["answers"]["_title_ok"] = True
        s["answers"].pop("_need_typed_title", None)

    try:
        q_result = generate_question(s["file_text"], s["answers"], s["questions_asked"], s.get("analysis"))
    except Exception:
        q_result = {}
    if q_result.get("ready"):
        s["ready"] = True
        return {"ready": True, "question": "", "options": []}

    s["_last_qtype"] = q_result.get("type", "")
    return {"question": q_result.get("question", ""), "options": q_result.get("options", []), "context": q_result.get("context", ""), "type": q_result.get("type", "")}

def strip_reasoning(text):
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if _heading_text(line):
            return "\n".join(lines[i:])
    return text

def _heading_text(line):
    stripped = line.strip()
    if not stripped or stripped.startswith("```") or stripped.startswith("==="):
        return None

    m = re.match(r'^(#{1,3})(\s*)(.*)', stripped)
    if m and (m.group(2) or not m.group(3) or len(m.group(1)) >= 2):
        return m.group(3).strip() or None

    for prefix in ["Abstract","Introduction","Literature","Related Work","Methodology",
                   "System Design","Implementation","Experimental","Results","Discussion",
                   "Conclusion","Future Work","References"]:
        clean = re.sub(r'^#+\s*', '', stripped)
        if clean.lower().startswith(prefix.lower()) and (clean.endswith(":") or len(clean) < len(prefix) + 5):
            return prefix

    return None

def parse_paper_text(paper_text, analysis, session_id):
    paper_text = strip_reasoning(paper_text)
    sections = []
    abstract = ""
    current_title = None
    current_content = []

    lines = paper_text.split("\n")
    md_title = None
    start_idx = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        if not stripped.startswith("##") and re.match(r'^#\s+\S', stripped):
            md_title = stripped.lstrip("#").strip()
            start_idx = i + 1
            break

    for line in lines[start_idx:]:
        stripped = line.strip()
        if not stripped or stripped.startswith("```") or stripped.startswith("==="):
            continue
        heading = _heading_text(line)
        if heading:
            content = "\n".join(current_content).strip()
            if current_title:
                sections.append({"title": current_title, "content": content})
            else:
                abstract = content
            current_title = heading
            current_content = []
        else:
            current_content.append(line)

    if current_title:
        sections.append({"title": current_title, "content": "\n".join(current_content).strip()})
    else:
        abstract = "\n".join(current_content).strip()

    title = md_title or analysis.get("title", "Research Paper")

    abstract_section = None
    other_sections = []
    for sec in sections:
        if sec["title"].lower().startswith("abstract") or (sec["title"].lower() in ("abstract", "abstract.")):
            abstract_section = sec["content"]
        else:
            other_sections.append(sec)

    if not abstract_section and other_sections:
        intro = other_sections[0]["content"]
        abstract_section = intro[:500]

    authors = analysis.get("authors") or ["Author A", "Author B"]
    affiliation = analysis.get("affiliation") or analysis.get("domain") or "Engineering"
    contact_email = analysis.get("contact_email", "")
    authors_data = []
    for i, author in enumerate(authors):
        authors_data.append({"name": author, "affiliation": affiliation, "email": contact_email})

    refs = []
    for sec in other_sections[:]:
        if sec["title"].lower().startswith("reference"):
            raw = sec["content"]
            parts = re.split(r'\n\s*(?=\[\d+\])', raw.strip())
            for p in parts:
                p = p.strip()
                if p and re.match(r'\[\d+\]', p):
                    refs.append({"citation": re.sub(r'^\[\d+\]\s*', '', p)})
            other_sections.remove(sec)
    if not refs:
        refs = [
            {"citation": f"J. Smith, ``Advances in {analysis.get('domain', 'Technology')},'' IEEE Trans., vol. 45, no. 3, pp. 123-135, 2023."},
            {"citation": f"B. Johnson, ``Modern {analysis.get('domain', 'Technology')} Systems,'' IEEE Conf. Proc., pp. 456-467, 2022."},
            {"citation": f"C. Williams, ``{analysis.get('domain', 'Technology')} Innovation,'' IEEE J., vol. 12, no. 4, pp. 789-801, 2023."},
        ]

    html_content = generate_ieee_html(
        title=title,
        authors=authors_data,
        abstract=abstract_section or abstract or paper_text[:500],
        sections=other_sections,
        keywords=analysis.get("keywords", [analysis.get("domain", "Technology")]),
        domain=analysis.get("domain", "Technology"),
        references=refs,
        session_id=session_id,
    )
    paper_json = {
        "title": title,
        "authors": authors_data,
        "abstract": abstract_section or abstract or "",
        "keywords": analysis.get("keywords", [analysis.get("domain", "Technology")]),
        "sections": other_sections,
        "references": [r["citation"] for r in refs],
    }
    base_name = analysis.get('title', 'paper').replace(' ', '_')
    return {
        "paper_text": paper_text,
        "html_content": html_content,
        "download_html": f"/api/download/{session_id}/html",
        "filename_html": f"{base_name}.html",
        "download_docx": f"/api/download/{session_id}/docx",
        "filename_docx": f"{base_name}.docx",
        "paper_json": paper_json,
    }

@app.get("/api/generate-paper-stream/{session_id}")
def handle_generate_paper_stream(session_id: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    def event_stream():
        paper_text = ""
        try:
            images_info = ""
            if s.get("images"):
                img_desc = []
                for img in s["images"]:
                    ctx = img.get("context", "")
                    ctx_snippet = (" — near text: " + ctx[:120].replace("\n", " ")) if ctx else ""
                    img_desc.append(f"  - {img['filename']} (page {img['page']}, {img['width']}x{img['height']}){ctx_snippet}")
                images_info = "\nExtracted images from document (you can reference them with ![caption](filename) markdown):\n" + "\n".join(img_desc)
            for token in generate_paper_stream(s["file_text"], s["answers"], s["analysis"], images_info):
                paper_text += token
                safe = token.replace("\n", "\\n").replace("\r", "\\r")
                yield f"data: {json.dumps({'type': 'token', 'content': safe})}\n\n"

            s["paper_text"] = paper_text
            result = parse_paper_text(paper_text, s["analysis"], session_id)
            s["html_content"] = result["html_content"]
            s["paper_json"] = result.get("paper_json")
            yield f"data: {json.dumps({'type': 'done', 'result': result})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")

def generate_pdf_from_html(html_content: str, paper_format: Optional[PaperFormat] = None, images: Optional[list] = None, session_dir: str = "") -> bytes:
    from fpdf import FPDF

    fmt = paper_format or PROCOMM

    def extract_section(regex, html, group=1):
        m = re.search(regex, html, re.DOTALL)
        return m.group(group).strip() if m else ""

    title = extract_section(r'<h1>(.*?)</h1>', html_content)
    abstract = extract_section(r'abstract-label">(.*?)</span>\s*<p>(.*?)</p>', html_content, 2)
    keywords = extract_section(r'kw-label">.*?</span>(.*?)</div>', html_content)
    abstract = re.sub(r'<[^>]+>', '', abstract)
    keywords = re.sub(r'<[^>]+>', '', keywords)

    author_blocks = re.findall(r'<div class="author">(.*?)</div>', html_content, re.DOTALL)
    authors = []
    for ab in author_blocks:
        name = re.sub(r'<[^>]+>', '', re.split(r'<br>', ab)[0]).strip()
        affil = re.sub(r'<[^>]+>', '', extract_section(r'class="affil">(.*?)</span>', ab)).strip()
        email = re.sub(r'<[^>]+>', '', extract_section(r'class="email">(.*?)</span>', ab)).strip()
        if name:
            authors.append((name, affil, email))

    sections_raw = re.findall(r'<h2>(.*?)</h2>(.*?)(?=<h2>|<div class="references"|$)', html_content, re.DOTALL)
    sections = []
    for t, body in sections_raw:
        t_clean = re.sub(r'<[^>]+>', '', t).strip()
        ps = re.findall(r'<p>(.*?)</p>', body, re.DOTALL)
        content = '\n\n'.join(re.sub(r'<[^>]+>', '', p).strip().replace('<br>', '\n') for p in ps if p.strip())
        # Restore image references from figure divs back into content
        figures = re.findall(r'<div class="figure">.*?<img src="[^"]*/([^"/]+)"[^>]*>.*?<div class="figcaption">(.*?)</div>', body, re.DOTALL)
        for fname, alt in figures:
            content += f'\n\n![{alt}]({fname})'
        tables = []
        for tbl in re.findall(r'<table class="ieee-table">(.*?)</table>', body, re.DOTALL):
            heads = [re.sub(r'<[^>]+>', '', h).strip() for h in re.findall(r'<th class="tablehead">(.*?)</th>', tbl)]
            rows = []
            for tr in re.findall(r'<tr>(.*?)</tr>', tbl, re.DOTALL):
                cells = [re.sub(r'<[^>]+>', '', c).strip() for c in re.findall(r'<td class="tabletext">(.*?)</td>', tr)]
                if cells:
                    rows.append(cells)
            if heads or rows:
                tables.append((heads, rows))
        if t_clean and (content or tables):
            sections.append((t_clean, content, tables))

    def ascii_safe(t):
        import unicodedata
        t = t.replace('\u2014', '--').replace('\u2013', '-')
        t = t.replace('\u2018', "'").replace('\u2019', "'")
        t = t.replace('\u201c', '"').replace('\u201d', '"')
        t = t.replace('\u00b2', '^2').replace('\u00b3', '^3')
        t = t.replace('\u00b0', ' deg ')
        t = t.replace('\u00b1', '+/-')
        t = t.replace('\u00d7', 'x').replace('\u00f7', '/')
        t = t.replace('\u03b1', 'alpha').replace('\u03b2', 'beta')
        t = t.replace('\u03b3', 'gamma').replace('\u03b4', 'delta')
        t = t.replace('\u03b8', 'theta').replace('\u03bb', 'lambda')
        t = t.replace('\u03bc', 'mu').replace('\u03c0', 'pi')
        t = t.replace('\u03c3', 'sigma').replace('\u03c9', 'omega')
        t = t.replace('\u0394', 'Delta').replace('\u03a3', 'Sigma')
        t = t.replace('\u2202', 'd').replace('\u2207', 'nabla')
        t = t.replace('\u2211', 'sum').replace('\u221e', 'inf')
        t = t.replace('\u2264', '<=').replace('\u2265', '>=')
        t = t.replace('\u2260', '!=')
        t = unicodedata.normalize('NFKD', t)
        return t.encode('ascii', 'replace').decode('ascii')
    sections = [(ascii_safe(t), ascii_safe(c), tables) for t, c, tables in sections]
    title = ascii_safe(title)
    abstract = ascii_safe(abstract)
    keywords = ascii_safe(keywords)
    authors = [(ascii_safe(n), ascii_safe(a), ascii_safe(e)) for n, a, e in authors]

    pw, _ = fmt.page_size_mm
    ml, mr, mt, mb = fmt.margins_mm
    content_w = pw - ml - mr
    col_gap = fmt.column_gap_mm
    col_w = (content_w - col_gap) / 2

    class IEEE_PDF(FPDF):
        def __init__(self):
            super().__init__(orientation="P", unit="mm", format="A4")
            self.col = 0
            self.col_x = [ml, ml + col_w + col_gap]
            self.col_w = col_w
            self.mt = mt
            self.b_margin = mb
            self.fmt = fmt
            self.header_drawn = False
            self.set_margins(ml, mt, mr)
            self.set_auto_page_break(auto=True, margin=mb)

        def header(self):
            if self.fmt.header_text and self.page_no() > 1:
                self.set_font(self.fmt.font_family, "", self.fmt.header_font_size)
                self.set_xy(ml, mt - 10)
                self.cell(content_w, 5, self.fmt.header_text, align="C", new_x="LMARGIN", new_y="NEXT")

        def footer(self):
            if self.fmt.footer_text:
                self.set_font(self.fmt.font_family, "", self.fmt.footer_font_size)
                self.set_xy(ml, self.h - 12)
                self.cell(content_w, 5, self.fmt.footer_text, align="C")

        def col_x_pos(self):
            return self.col_x[self.col]

        def col_ln(self, h=None):
            if h is not None:
                self.y += h
            self.x = self.col_x_pos()
            return self

        def _perform_page_break(self):
            if self.col == 0:
                self.col = 1
                self.x = self.col_x[1]
                by = getattr(self, 'body_start_y', None)
                if by is not None and self.page_no() == 1:
                    self.y = by
                else:
                    self.y = self.mt + (6 if self.fmt.header_text else 0)
            else:
                super()._perform_page_break()
                self.x = self.col_x[0]
                self.y = self.mt + (6 if self.fmt.header_text else 0)
                self.col = 0

        def ensure_col(self, h):
            if self.will_page_break(h):
                self._perform_page_break()

        def draw_smallcaps(self, text, size, line_h, align="C"):
            x = self.x
            y = self.y
            w = self.col_w
            def cwidth(ch):
                self.set_font(self.fmt.font_family, "", int(size * 0.72) if ch.islower() else size)
                return self.get_string_width(ch.upper())
            words = text.split()
            lines, cur, curw = [], [], 0
            sp = self.get_string_width(" ")
            for word in words:
                ww = sum(cwidth(ch) for ch in word)
                if cur and curw + sp + ww > w:
                    lines.append(" ".join(cur))
                    cur, curw = [word], ww
                else:
                    if cur:
                        curw += sp
                    cur.append(word)
                    curw += ww
            if cur:
                lines.append(" ".join(cur))
            for li, line in enumerate(lines):
                yy = y + li * line_h
                total = sum(cwidth(ch) if ch != " " else sp for ch in line)
                sx = x + (w - total) / 2 if align == "C" else x
                self.set_xy(sx, yy)
                for ch in line:
                    if ch == " ":
                        self.set_font(self.fmt.font_family, "", size)
                        self.cell(sp, line_h, " ", new_x="RIGHT", new_y="TOP")
                        continue
                    self.set_font(self.fmt.font_family, "", int(size * 0.72) if ch.islower() else size)
                    wc = self.get_string_width(ch.upper())
                    self.cell(wc, line_h, ch.upper(), new_x="RIGHT", new_y="TOP")
            self.set_font(self.fmt.font_family, "", size)
            self.set_xy(x, y + len(lines) * line_h)
            return len(lines)

        def draw_section(self, sec_title, size, line_h):
            if self.fmt.section_centered:
                n = self.draw_smallcaps(sec_title, size, line_h, "C")
                self.col_ln(n * line_h + 2)
            else:
                self.set_font(self.fmt.font_family, "B", size)
                self.multi_cell(self.col_w, line_h * 0.7, sec_title, align="L", new_x="LEFT", new_y="NEXT")
                self.x = self.col_x_pos()
                self.col_ln(2)

    pdf = IEEE_PDF()
    pdf.add_page()
    header_h = 6 if fmt.header_text else 0
    cy = mt + header_h

    # --- Title block (full width, single column) ---
    pdf.set_font(fmt.font_family, "B", fmt.title_font_size)
    pdf.set_xy(ml, cy)
    pdf.multi_cell(content_w, 9, title, align="C", new_x="LMARGIN", new_y="NEXT")
    cy = pdf.get_y() + 3
    for name, affil, email in authors:
        pdf.set_xy(ml, cy)
        pdf.set_font(fmt.font_family, "", fmt.author_font_size)
        pdf.multi_cell(content_w, 6, name, align="C", new_x="LMARGIN", new_y="NEXT")
        cy = pdf.get_y()
        if affil:
            pdf.set_xy(ml, cy)
            pdf.set_font(fmt.font_family, "I" if fmt.abstract_bold_italic else "", 10)
            pdf.multi_cell(content_w, 5, affil, align="C", new_x="LMARGIN", new_y="NEXT")
            cy = pdf.get_y()
        if email:
            pdf.set_xy(ml, cy)
            pdf.set_font(fmt.font_family, "", 10)
            pdf.multi_cell(content_w, 5, email, align="C", new_x="LMARGIN", new_y="NEXT")
            cy = pdf.get_y()
    cy += 3
    if abstract:
        pdf.set_xy(ml, cy)
        if fmt.abstract_bold_italic:
            pdf.set_font(fmt.font_family, "BI", 10)
            pdf.multi_cell(content_w, 5, f"Abstract - {abstract}", align="J", new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font(fmt.font_family, "B", 10)
            pdf.cell(content_w, 5, "Abstract", align="L", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font(fmt.font_family, "", fmt.body_font_size)
            pdf.multi_cell(content_w, 5, abstract, align="J", new_x="LMARGIN", new_y="NEXT")
        cy = pdf.get_y() + 2
    if keywords:
        pdf.set_xy(ml, cy)
        pdf.set_font(fmt.font_family, "I", 10)
        pdf.multi_cell(content_w, 5, f"Index Terms - {keywords}", align="J", new_x="LMARGIN", new_y="NEXT")
        cy = pdf.get_y() + 4

    pdf.set_xy(pdf.col_x[0], cy)
    pdf.col = 0
    pdf.body_start_y = cy

    # --- Body (two columns) ---
    for sec_title, sec_content, tables in sections:
        size, line_h = fmt.body_font_size, fmt.body_font_size + 2
        hdr_h = 16
        pdf.ensure_col(hdr_h)
        pdf.draw_section(sec_title, size, line_h)

        img_re = re.compile(r'!\[(.*?)\]\(([^)]+)\)')
        for para in sec_content.split('\n'):
            para = para.strip()
            if not para:
                continue
            img_m = img_re.match(para)
            if img_m:
                fname = os.path.basename(img_m.group(2))
                alt = img_m.group(1)
                img_path = os.path.join(session_dir, "images", fname) if session_dir else ""
                if img_path and os.path.isfile(img_path):
                    try:
                        from PIL import Image as PILImage
                        with PILImage.open(img_path) as pim:
                            iw, ih = pim.size
                        pw_mm, _ = fmt.page_size_mm
                        max_w = pdf.col_w - 4
                        scale = min(max_w, pw_mm * 0.4) / iw if iw > 0 else 0.5
                        img_h_mm = ih * scale
                        total_h = img_h_mm + 10
                        pdf.ensure_col(total_h)
                        pdf.image(img_path, x=pdf.x + 2, w=min(max_w, pw_mm * 0.4))
                        pdf.col_ln(max(img_h_mm * 0.15, 3))
                        pdf.set_font(fmt.font_family, "", fmt.body_font_size)
                        pdf.multi_cell(pdf.col_w, 4, alt, align="C", new_x="LEFT", new_y="NEXT")
                        pdf.x = pdf.col_x_pos()
                        pdf.col_ln(2)
                    except Exception:
                        pdf.set_font(fmt.font_family, "", fmt.body_font_size)
                        pdf.multi_cell(pdf.col_w, 4, f"[Image: {alt}]", align="C", new_x="LEFT", new_y="NEXT")
                        pdf.x = pdf.col_x_pos()
                continue
            pdf.set_font(fmt.font_family, "", fmt.body_font_size)
            pdf.multi_cell(pdf.col_w, 5, para, align="J", new_x="LEFT", new_y="NEXT")
            pdf.x = pdf.col_x_pos()

        for heads, rows in tables:
            table_rows = ([heads] if heads else []) + rows
            for idx, row in enumerate(table_rows):
                line = " | ".join(row)
                pdf.set_font(fmt.font_family, "B" if (heads and idx == 0) else "", 9)
                pdf.multi_cell(pdf.col_w, 5, line, align="L", new_x="LEFT", new_y="NEXT")
                pdf.x = pdf.col_x_pos()
            pdf.col_ln(2)

    return bytes(pdf.output())


# ─── DOCX generator (IEEE) ────────────────────────────────────────────────────


def _set_continuous_section_break(doc):
    """After doc.add_section(), find the embedded sectPr in the previous
    section's last paragraph and change its break type to continuous."""
    body = doc.element.body
    for p in reversed(body.findall(qn('w:p'))):
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            embedded = pPr.find(qn('w:sectPr'))
            if embedded is not None:
                type_elem = embedded.find(qn('w:type'))
                if type_elem is None:
                    type_elem = OxmlElement('w:type')
                    embedded.insert(0, type_elem)
                type_elem.set(qn('w:val'), 'continuous')
                return True
    return False


def _add_run(paragraph, text, size=10, bold=False, italic=False, font_name='Times New Roman', small_caps=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if small_caps:
        run.font.small_caps = True
    return run


def _set_cols(sectPr, num=None, space='720'):
    """Set or update the w:cols element on a sectPr.

    Removes any existing w:cols, then creates a new one.
    num=None → single column (attribute omitted).
    num='2'  → two columns with specified gutter space.
    """
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    existing = sectPr.find(ns + 'cols')
    if existing is not None:
        sectPr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:space'), space)
    if num is not None:
        cols.set(qn('w:num'), num)
    sectPr.append(cols)


def _remove_table_borders(table):
    """Remove all borders from a python-docx table via XML manipulation."""
    tbl_pr = table._tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tbl_pr.append(borders)


def _is_table_sep(s):
    """Detect markdown table separator row like |---|---|."""
    s = s.strip().strip("|")
    return bool(s) and set(s) <= set("-: |") and "-" in s


def _is_caption(s):
    """Detect figure/table caption."""
    return bool(re.match(r'^(figure|table)\s*\d*\.?\s', s, re.I))


def _render_body_docx(doc, content, session_dir=""):
    """Parse body content text and add paragraphs, tables, and images to the document."""
    img_re = re.compile(r'!\[(.*?)\]\(([^)]+)\)')
    lines = [l.rstrip() for l in content.split("\n")]
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # ---- Image ----
        img_m = img_re.match(stripped)
        if img_m:
            fname = os.path.basename(img_m.group(2))
            alt = img_m.group(1)
            img_path = os.path.join(session_dir, "images", fname) if session_dir else ""
            if img_path and os.path.isfile(img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(3))
                    p2 = doc.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run(p2, alt, size=9, italic=True)
                except Exception:
                    pass
            i += 1
            continue

        # ---- Caption ----
        if _is_caption(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0
            _add_run(p, stripped, size=10, small_caps=True)
            i += 1
            continue

        # ---- Markdown table ----
        if "|" in stripped and i + 1 < n and _is_table_sep(lines[i + 1]):
            rows = []
            while i < n and "|" in lines[i].strip():
                rows.append(lines[i].strip())
                i += 1
            header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
            data = []
            for r in rows[2:]:
                if "|" in r:
                    data.append([c.strip() for c in r.strip().strip("|").split("|")])
            if data:
                tbl = doc.add_table(rows=len(data) + 1, cols=len(header))
                tbl.style = 'Table Grid'
                for j, h in enumerate(header):
                    cell = tbl.cell(0, j)
                    cp = cell.paragraphs[0]
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run(cp, h, size=9, bold=True)
                for ri, row in enumerate(data):
                    for ci, val in enumerate(row):
                        cell = tbl.cell(ri + 1, ci)
                        cp = cell.paragraphs[0]
                        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        _add_run(cp, val, size=9)
            continue

        # ---- Regular paragraph ----
        para = [stripped]
        i += 1
        while i < n and lines[i].strip() and "|" not in lines[i] and not _is_table_sep(lines[i]) and not _is_caption(lines[i].strip()):
            para.append(lines[i].strip())
            i += 1

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.17)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        _add_run(p, ' '.join(para), size=10)


def generate_ieee_docx(paper_json: dict, session_dir: str = "") -> bytes:
    """Generate a 3-section IEEE-format DOCX with continuous section breaks.

    Section 1 — single column: title, author table, abstract, keywords.
    Section 2 — two columns: body sections (Introduction → Acknowledgements).
    Section 3 — two columns: references.
    """
    doc = Document()

    # ── Page setup ──
    section1 = doc.sections[0]
    section1.top_margin = Inches(0.75)
    section1.bottom_margin = Inches(0.75)
    section1.left_margin = Inches(0.65)
    section1.right_margin = Inches(0.65)

    # ── Default style (Normal) ──
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.line_spacing = 1.0

    # ═══════════════════════════════════════════════════════════════
    # SECTION 1 — Single column, full width
    # ═══════════════════════════════════════════════════════════════
    sect_pr_1 = section1._sectPr
    _set_cols(sect_pr_1, num=None, space='720')   # single column
    title_pg = OxmlElement('w:titlePg')            # first-page different footer
    sect_pr_1.append(title_pg)

    # ── Title / Author table (3 columns, borderless) ──
    tbl_widths = [3271, 3273, 3276]          # dxa (twips)
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    # Set gridCol widths in twips (dxa)
    tbl_grid = table._tbl.find(qn('w:tblGrid'))
    if tbl_grid is not None:
        gcs = tbl_grid.findall(qn('w:gridCol'))
        for i, w in enumerate(tbl_widths):
            gcs[i].set(qn('w:w'), str(w))

    # Row 1 — Title merged across all 3 columns
    title_cell = table.cell(0, 0)
    title_cell.merge(table.cell(0, 2))
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    p_title.clear()
    _add_run(p_title, paper_json.get('title', ''), size=24, bold=True)

    # Row 2 — Authors in middle column
    authors = paper_json.get('authors', [])
    author_cell = table.cell(1, 1)
    p_author = author_cell.paragraphs[0]
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(12)
    p_author.clear()
    author_lines = []
    for a in authors:
        parts = [a.get('name', '')]
        affil = a.get('affiliation', '')
        if affil:
            parts.append(affil)
        email = a.get('email', '')
        if email:
            parts.append(email)
        author_lines.append(' — '.join(parts))
    _add_run(p_author, '\n'.join(author_lines), size=12)

    # ── Abstract (bold + italic, justified) ──
    abstract_text = paper_json.get('abstract', '')
    if abstract_text:
        p_abs = doc.add_paragraph()
        p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_abs.paragraph_format.space_before = Pt(0)
        p_abs.paragraph_format.space_after = Pt(0)
        p_abs.paragraph_format.line_spacing = Pt(12)   # 240 twips
        _add_run(p_abs, 'Abstract—', size=10, bold=True, italic=True)
        _add_run(p_abs, abstract_text, size=10, bold=True, italic=True)

    # ── Index Terms (italic) ──
    keywords = paper_json.get('keywords', [])
    if keywords:
        p_kw = doc.add_paragraph()
        p_kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_kw.paragraph_format.space_before = Pt(0)
        p_kw.paragraph_format.space_after = Pt(12)
        p_kw.paragraph_format.line_spacing = Pt(12)
        _add_run(p_kw, 'Index Terms—', size=10, bold=True, italic=True)
        _add_run(p_kw, ', '.join(keywords), size=10, italic=True)

    # ═══════════════════════════════════════════════════════════════
    # CONTINUOUS BREAK → SECTION 2  (two columns)
    # ═══════════════════════════════════════════════════════════════
    doc.add_section()
    _set_continuous_section_break(doc)

    sect_pr_2 = doc.sections[1]._sectPr
    _set_cols(sect_pr_2, num='2', space='461')

    # ── Section 2: Body ──
    sections_data = paper_json.get('sections', [])
    for sec in sections_data:
        sec_title = sec.get('title', '').strip()
        sec_content = sec.get('content', '').strip()

        if not sec_title:
            continue

        # Detect subsection: title starts with number+period (e.g. "2.1")
        is_subsection = bool(re.match(r'^(\d+\.\d+)', sec_title))

        if is_subsection:
            # Heading 2 — italic, left-aligned
            p_h2 = doc.add_paragraph()
            p_h2.paragraph_format.space_before = Pt(9)
            p_h2.paragraph_format.space_after = Pt(3)
            p_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_run(p_h2, sec_title, size=10, italic=True)
        else:
            # Heading 1 — centered, small caps
            p_h1 = doc.add_paragraph()
            p_h1.paragraph_format.space_before = Pt(12)
            p_h1.paragraph_format.space_after = Pt(6)
            p_h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_h1, sec_title, size=10, bold=True, small_caps=True)

        # Body paragraphs
        if sec_content:
            _render_body_docx(doc, sec_content, session_dir)

    # ═══════════════════════════════════════════════════════════════
    # CONTINUOUS BREAK → SECTION 3  (two columns, references)
    # ═══════════════════════════════════════════════════════════════
    doc.add_section()
    _set_continuous_section_break(doc)

    sect_pr_3 = doc.sections[2]._sectPr
    _set_cols(sect_pr_3, num='2', space='461')

    # ── Section 3: References ──
    refs = paper_json.get('references', [])
    if refs:
        p_ref_h = doc.add_paragraph()
        p_ref_h.paragraph_format.space_before = Pt(12)
        p_ref_h.paragraph_format.space_after = Pt(6)
        p_ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_ref_h, 'References', size=10, bold=True, small_caps=True)

        for i, ref in enumerate(refs, 1):
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.space_before = Pt(0)
            p_ref.paragraph_format.space_after = Pt(6)
            p_ref.paragraph_format.line_spacing = 1.0
            p_ref.paragraph_format.left_indent = Inches(0.25)
            p_ref.paragraph_format.first_line_indent = Inches(-0.25)
            _add_run(p_ref, f'[{i}] {ref}', size=10)

    # ── Save to bytes ──
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_markdown(title, abstract, sections, keywords, refs):
    parts = [f"# {title}", f"## Abstract\n{abstract}"]
    for sec in sections:
        parts.append(f"## {sec.get('title', '')}\n{sec.get('content', '')}")
    ref_lines = ["## References"]
    for i, r in enumerate(refs, 1):
        r = (r or "").strip()
        ref_lines.append(r if r.startswith("[") else f"[{i}] {r}")
    return "\n\n".join(parts + ref_lines)


# ─── Resume helpers ───────────────────────────────────────────────────────────

def generate_resume_html(resume_data: dict) -> str:
    """Render a professional single-column resume HTML."""
    name = resume_data.get("name", "Your Name")
    email = resume_data.get("email", "")
    phone = resume_data.get("phone", "")
    linkedin = resume_data.get("linkedin", "")
    github = resume_data.get("github", "")
    location = resume_data.get("location", "")

    contact_parts = []
    if location:
        contact_parts.append(f'<span class="ci-icon">&#9906;</span> {location}')
    if email:
        contact_parts.append(f'<span class="ci-icon">&#9993;</span> {email}')
    if phone:
        contact_parts.append(f'<span class="ci-icon">&#9742;</span> {phone}')
    if linkedin:
        contact_parts.append(f'<span class="ci-icon">in</span> {linkedin}')
    if github:
        contact_parts.append(f'<span class="ci-icon">&#9733;</span> {github}')
    contact_str = '&nbsp;&nbsp;&nbsp;'.join(contact_parts)

    sections_html = ""
    section_order = ["Summary", "Education", "Experience", "Skills", "Projects", "Certifications"]
    for section_name in section_order:
        items = resume_data.get(section_name.lower(), [])
        if not items:
            continue
        sections_html += f'<div class="section"><div class="section-title">{section_name}</div>'
        if section_name == "Summary":
            text = items[0] if isinstance(items, list) else items
            sections_html += f'<p class="summary-text">{text}</p>'
        elif section_name == "Skills":
            sections_html += '<div class="skills">'
            for skill in items:
                sections_html += f'<span class="skill">{skill}</span>'
            sections_html += '</div>'
        elif section_name in ("Education", "Certifications"):
            for item in items:
                if isinstance(item, dict):
                    title = item.get("degree") or item.get("name") or item.get("text", "")
                    school = item.get("school", "")
                    year = item.get("year", "")
                    gpa = item.get("gpa", "")
                    sections_html += '<div class="item">'
                    sections_html += '<div class="item-header">'
                    sections_html += f'<div class="item-title">{title}</div>'
                    if year:
                        sections_html += f'<div class="item-date">{year}</div>'
                    sections_html += '</div>'
                    if school:
                        sections_html += f'<div class="item-subtitle">{school}</div>'
                    if gpa:
                        sections_html += f'<div class="item-detail">GPA: {gpa}</div>'
                    sections_html += '</div>'
                else:
                    sections_html += f'<div class="item"><p>{item}</p></div>'
        elif section_name == "Experience":
            for item in items:
                if isinstance(item, dict):
                    title = item.get("role") or item.get("title") or ""
                    company = item.get("company", "")
                    dates = item.get("dates", "")
                    location = item.get("location", "")
                    sections_html += '<div class="item">'
                    sections_html += '<div class="item-header">'
                    sections_html += f'<div class="item-title">{title}</div>'
                    if dates:
                        sections_html += f'<div class="item-date">{dates}</div>'
                    sections_html += '</div>'
                    if company:
                        loc_str = f" &middot; {location}" if location else ""
                        sections_html += f'<div class="item-subtitle">{company}{loc_str}</div>'
                    bullets = item.get("bullets", [])
                    if bullets:
                        sections_html += '<ul class="bullets">'
                        for b in bullets:
                            sections_html += f'<li>{b}</li>'
                        sections_html += '</ul>'
                    sections_html += '</div>'
                else:
                    sections_html += f'<div class="item"><p>{item}</p></div>'
        elif section_name == "Projects":
            for item in items:
                if isinstance(item, dict):
                    title = item.get("name") or item.get("title") or ""
                    tech = item.get("tech", "")
                    link = item.get("link", "")
                    sections_html += '<div class="item">'
                    sections_html += '<div class="item-header">'
                    sections_html += f'<div class="item-title">{title}</div>'
                    if tech:
                        sections_html += f'<div class="item-tech">{tech}</div>'
                    sections_html += '</div>'
                    if link:
                        sections_html += f'<div class="item-link"><a href="{link}">{link}</a></div>'
                    desc = item.get("description", "")
                    if desc:
                        sections_html += f'<p class="project-desc">{desc}</p>'
                    for b in item.get("bullets", []):
                        sections_html += f'<p class="project-bullet">&bull; {b}</p>'
                    sections_html += '</div>'
                else:
                    sections_html += f'<div class="item"><p>{item}</p></div>'
        sections_html += '</div>'

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>{name} - Resume</title>
<style>
  @page {{ size: letter; margin: 0.45in 0.5in; }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: "Helvetica Neue", Helvetica, Arial, sans-serif; font-size: 10pt; color: #2d2d2d; line-height: 1.4; }}
  .resume {{ max-width: 7.5in; margin: 0 auto; padding: 0.35in 0.45in; }}
  .header {{ text-align: center; padding-bottom: 10px; margin-bottom: 14px; }}
  .header::after {{ content: ""; display: block; height: 2px; background: linear-gradient(to right, transparent, #1a5276, transparent); margin-top: 10px; }}
  .name {{ font-size: 24pt; font-weight: 700; color: #1a1a2e; letter-spacing: 1px; margin-bottom: 3px; }}
  .contact {{ font-size: 8.5pt; color: #555; line-height: 1.6; }}
  .ci-icon {{ font-weight: bold; color: #1a5276; margin-right: 1px; }}
  .section {{ margin-bottom: 12px; }}
  .section-title {{ font-size: 11pt; font-weight: 700; color: #1a5276; text-transform: uppercase; letter-spacing: 1px; border-bottom: 1.5px solid #2980b9; padding-bottom: 2px; margin-bottom: 7px; }}
  .summary-text {{ font-size: 9.5pt; color: #444; line-height: 1.5; }}
  .item {{ margin-bottom: 9px; }}
  .item-header {{ display: flex; justify-content: space-between; align-items: baseline; }}
  .item-title {{ font-weight: 700; font-size: 10pt; color: #1a1a2e; }}
  .item-date {{ color: #666; font-size: 9pt; white-space: nowrap; }}
  .item-subtitle {{ font-style: italic; color: #555; font-size: 9.5pt; margin-bottom: 2px; }}
  .item-detail {{ font-size: 9pt; color: #555; margin-top: 1px; }}
  .item-tech {{ font-size: 9pt; color: #2980b9; }}
  .item-link a {{ font-size: 8.5pt; color: #2980b9; text-decoration: none; }}
  .item-link a:hover {{ text-decoration: underline; }}
  ul.bullets {{ margin: 2px 0 0 14px; padding: 0; }}
  ul.bullets li {{ font-size: 9.5pt; color: #444; margin-bottom: 2px; line-height: 1.45; }}
  .project-desc {{ font-size: 9.5pt; color: #444; margin-top: 1px; }}
  .project-bullet {{ font-size: 9pt; color: #555; margin-left: 12px; }}
  .skills {{ display: flex; flex-wrap: wrap; gap: 4px 6px; }}
  .skill {{ background: #eaf2f8; color: #1a5276; padding: 2px 10px; border-radius: 2px; font-size: 9pt; border: 1px solid #d4e6f1; }}
  @media print {{ body {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }} .resume {{ padding: 0; }} .item-link a {{ text-decoration: none; }} }}
</style></head><body>
<div class="resume">
  <div class="header">
    <div class="name">{name}</div>
    <div class="contact">{contact_str}</div>
  </div>
  {sections_html}
</div>
</body></html>"""


def parse_resume_text(resume_text: str) -> dict:
    """Parse a markdown resume into structured JSON."""
    result = {
        "name": "",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "location": "",
        "summary": [],
        "education": [],
        "experience": [],
        "skills": [],
        "projects": [],
        "certifications": [],
    }
    lines = resume_text.strip().split("\n")
    current_section = None
    current_item = None
    section_map = {
        "summary": "summary",
        "education": "education",
        "experience": "experience",
        "skills": "skills",
        "projects": "projects",
        "certifications": "certifications",
    }

    def extract_email(s):
        m = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', s)
        return m.group(0) if m else ""

    def extract_phone(s):
        m = re.search(r'[\+]?[\d\-\(\)\s]{7,}', s)
        return m.group(0).strip() if m else ""

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("```"):
            continue

        # Try to extract contact info from first few lines before sections
        if not any(stripped.startswith("#") for s in ["## ", "### "]):
            if current_section is None:
                if not result["name"] and not stripped.startswith(("- ", "* ")):
                    if "linkedin" in stripped.lower():
                        result["linkedin"] = stripped.replace("- ", "").replace("* ", "").strip()
                    elif "github" in stripped.lower():
                        result["github"] = stripped.replace("- ", "").replace("* ", "").strip()
                    elif "@" in stripped:
                        result["email"] = extract_email(stripped) or stripped.strip()
                    elif re.search(r'[\+]?\d[\d\-\(\)\s]{7,}', stripped):
                        result["phone"] = extract_phone(stripped)
                    else:
                        # If no section started, this might be contact line
                        pass

        # H2 section heading
        if stripped.startswith("## "):
            heading = stripped[3:].strip().lower()
            current_section = section_map.get(heading)
            current_item = None
            continue

        if current_section is None:
            continue

        if current_section == "summary":
            result["summary"].append(stripped.lstrip("- ").lstrip("* ").strip())
        elif current_section == "skills":
            clean = stripped.lstrip("- ").lstrip("* ").strip()
            clean = clean.replace("**", "").replace("__", "").strip()
            if clean:
                result["skills"].append(clean)
        elif current_section == "education":
            if stripped.startswith("### ") or stripped.startswith("**"):
                raw = stripped.lstrip("#").strip().replace("**", "").replace("__", "").strip()
                # Try to parse "Degree — School, Year" or "Degree | School | Year"
                item = {"text": raw}
                for sep in [" — ", " | ", " - ", " – "]:
                    parts = raw.split(sep, 2)
                    if len(parts) >= 2:
                        item["degree"] = parts[0].strip()
                        remaining = parts[1].strip()
                        if len(parts) >= 3:
                            item["school"] = remaining
                            item["year"] = parts[2].strip()
                        else:
                            # Try splitting remaining on comma
                            if "," in remaining:
                                s, y = remaining.rsplit(",", 1)
                                yy = y.strip()
                                gpa_m = re.search(r'[Gg][Pp][Aa][:\s]*([\d.]+)', yy)
                                if gpa_m:
                                    item["gpa"] = gpa_m.group(1)
                                    no_gpa = remaining.rsplit(",", 1)[0].strip()
                                    if "," in no_gpa:
                                        s2, y2 = no_gpa.rsplit(",", 1)
                                        item["school"] = s2.strip()
                                        item["year"] = y2.strip()
                                    else:
                                        item["school"] = no_gpa
                                else:
                                    item["school"] = s.strip()
                                    item["year"] = yy
                            else:
                                item["school"] = remaining
                        break
                # Extract GPA if not already found
                if "gpa" in raw.lower() and not item.get("gpa"):
                    gpa_m = re.search(r'[Gg][Pp][Aa][:\s]*([\d.]+)', raw)
                    if gpa_m:
                        item["gpa"] = gpa_m.group(1)
                result["education"].append(item)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                clean = stripped[2:].strip().replace("**", "").replace("__", "")
                result["education"].append({"text": clean})
            else:
                result["education"].append({"text": stripped.replace("**", "").replace("__", "")})
        elif current_section == "certifications":
            if stripped.startswith("- ") or stripped.startswith("* "):
                clean = stripped[2:].strip().replace("**", "").replace("__", "")
                result["certifications"].append({"text": clean})
            elif stripped.startswith("**") or stripped.startswith("###"):
                clean = stripped.lstrip("#").strip().replace("**", "").replace("__", "").strip()
                result["certifications"].append({"text": clean})
            else:
                result["certifications"].append({"text": stripped.replace("**", "").replace("__", "")})
        elif current_section == "experience":
            if stripped.startswith("### "):
                raw = stripped[4:].strip().replace("**", "").replace("__", "").strip()
                # Parse "Role | Company | Location | Dates"
                item = {"title": raw, "bullets": []}
                parts = [p.strip() for p in raw.replace(" — ", " | ").replace(" – ", " | ").split("|")]
                if len(parts) >= 1:
                    item["role"] = parts[0].strip()
                if len(parts) >= 2:
                    item["company"] = parts[1].strip()
                if len(parts) >= 3:
                    # Could be location or dates
                    if any(kw in parts[2].lower() for kw in ("present", "20", "19")):
                        item["dates"] = parts[2].strip()
                    else:
                        item["location"] = parts[2].strip()
                if len(parts) >= 4:
                    item["dates"] = parts[3].strip()
                current_item = item
                result["experience"].append(current_item)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                bullet = stripped[2:].strip().replace("**", "").replace("__", "")
                if current_item:
                    current_item["bullets"].append(bullet)
                else:
                    current_item = {"title": "", "role": "", "bullets": [bullet]}
                    result["experience"].append(current_item)
            else:
                clean = stripped.replace("**", "").replace("__", "")
                if current_item:
                    current_item["bullets"].append(clean)
                else:
                    current_item = {"title": clean, "role": clean, "bullets": []}
                    result["experience"].append(current_item)
        elif current_section == "projects":
            if stripped.startswith("### ") or stripped.startswith("**"):
                raw = stripped.lstrip("#").strip().replace("**", "").replace("__", "").strip()
                # Parse "Name | Tech" link
                item = {"title": raw, "name": raw, "bullets": []}
                if " | " in raw:
                    parts = raw.split(" | ", 1)
                    item["title"] = parts[0].strip()
                    item["name"] = parts[0].strip()
                    if "tech" in parts[1].lower() or ":" in parts[1]:
                        item["tech"] = parts[1].strip()
                    else:
                        item["tech"] = parts[1].strip()
                current_item = item
                result["projects"].append(current_item)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                bullet = stripped[2:].strip().replace("**", "").replace("__", "")
                if current_item:
                    current_item["bullets"].append(bullet)
                else:
                    current_item = {"title": "", "name": "", "bullets": [bullet]}
                    result["projects"].append(current_item)
            else:
                clean = stripped.replace("**", "").replace("__", "")
                if current_item:
                    current_item["bullets"].append(clean)
                else:
                    current_item = {"title": clean, "name": clean, "bullets": []}
                    result["projects"].append(current_item)

    # Flatten summary
    result["summary"] = " ".join(result["summary"]) if result["summary"] else ""

    # Try to extract name from first line if not in section
    if not result["name"] and lines:
        first = lines[0].strip()
        if not first.startswith("#") and not first.startswith("##"):
            result["name"] = first

    return result


def generate_resume_pdf(resume_data: dict) -> bytes:
    """Generate a single-column resume PDF directly from resume_data dict."""
    from fpdf import FPDF

    def ascii_safe(t):
        if not isinstance(t, str):
            t = str(t)
        t = t.replace('\u2014', '--').replace('\u2013', '-')
        t = t.replace('\u2018', "'").replace('\u2019', "'")
        t = t.replace('\u201c', '"').replace('\u201d', '"')
        t = t.replace('\u00b2', '^2').replace('\u00b3', '^3')
        t = t.replace('\u00d7', 'x').replace('\u00f7', '/')
        t = t.replace('\u2022', '--').replace('\u2023', '-')
        return t.encode('ascii', 'replace').decode('ascii')

    name = ascii_safe(resume_data.get("name", "Resume"))
    email = ascii_safe(resume_data.get("email", ""))
    phone = ascii_safe(resume_data.get("phone", ""))
    linkedin = ascii_safe(resume_data.get("linkedin", ""))
    github_acc = ascii_safe(resume_data.get("github", ""))
    loc = ascii_safe(resume_data.get("location", ""))
    summary = ascii_safe(resume_data.get("summary", ""))
    education = resume_data.get("education", [])
    experience = resume_data.get("experience", [])
    skills = resume_data.get("skills", [])
    projects = resume_data.get("projects", [])
    certifications = resume_data.get("certifications", [])

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    ml = mr = 18
    pw = 210
    content_w = pw - ml - mr

    def section_header(title):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 82, 118)
        pdf.cell(content_w, 7, ascii_safe(title.upper()), new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(41, 128, 185)
        pdf.set_line_width(0.4)
        pdf.line(ml, pdf.get_y(), ml + content_w, pdf.get_y())
        pdf.ln(3)
        pdf.set_text_color(45, 45, 45)

    # Header
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_xy(ml, 18)
    pdf.cell(content_w, 10, name, align="C", new_x="LMARGIN", new_y="NEXT")

    contact_parts = [p for p in [loc, email, phone, linkedin, github_acc] if p]
    if contact_parts:
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_xy(ml, pdf.get_y() + 2)
        pdf.cell(content_w, 5, "  |  ".join(contact_parts), align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(3)
    pdf.set_draw_color(26, 82, 118)
    pdf.set_line_width(0.5)
    pdf.line(ml, pdf.get_y(), ml + content_w, pdf.get_y())
    pdf.ln(5)

    # Summary
    if summary:
        section_header("Summary")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_x(ml)
        pdf.multi_cell(content_w, 5, ascii_safe(summary), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # Education
    if education:
        section_header("Education")
        for item in education:
            if isinstance(item, dict):
                text = ascii_safe(item.get("degree") or item.get("name") or item.get("text", ""))
                school = ascii_safe(item.get("school", ""))
                year = ascii_safe(item.get("year", ""))
                gpa = ascii_safe(item.get("gpa", ""))
                if text:
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.set_x(ml)
                    if year:
                        pdf.cell(content_w - 30, 5, text, new_x="LEFT", new_y="TOP")
                        pdf.set_font("Helvetica", "I", 9)
                        pdf.cell(30, 5, year, align="R", new_x="LMARGIN", new_y="NEXT")
                    else:
                        pdf.cell(content_w, 5, text, new_x="LMARGIN", new_y="NEXT")
                if school:
                    pdf.set_font("Helvetica", "I", 9)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 4, school, new_x="LMARGIN", new_y="NEXT")
                if gpa:
                    pdf.set_font("Helvetica", "", 9)
                    pdf.set_text_color(85, 85, 85)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 4, f"GPA: {gpa}", new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(45, 45, 45)
            else:
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_x(ml)
                pdf.cell(content_w, 5, ascii_safe(str(item)), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1.5)
        pdf.ln(2)

    # Experience
    if experience:
        section_header("Experience")
        for item in experience:
            if isinstance(item, dict):
                role = ascii_safe(item.get("role") or item.get("title", ""))
                company = ascii_safe(item.get("company", ""))
                dates = ascii_safe(item.get("dates", ""))
                exp_loc = ascii_safe(item.get("location", ""))
                if role:
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.set_x(ml)
                    pdf.cell(content_w - 40, 5, role, new_x="LEFT", new_y="TOP")
                    pdf.set_font("Helvetica", "I", 9)
                    pdf.cell(40, 5, dates, align="R", new_x="LMARGIN", new_y="NEXT")
                if company:
                    subtitle = company + (f" -- {exp_loc}" if exp_loc else "")
                    pdf.set_font("Helvetica", "I", 9)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 4, subtitle, new_x="LMARGIN", new_y="NEXT")
                for bullet in item.get("bullets", []):
                    clean = ascii_safe(re.sub(r'^[\u2022\-*\s]+', '', str(bullet)).strip())
                    if clean:
                        pdf.set_font("Helvetica", "", 9.5)
                        pdf.set_x(ml + 3)
                        pdf.multi_cell(content_w - 5, 4.5, f"- {clean}", new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.set_x(ml)
                pdf.multi_cell(content_w, 5, ascii_safe(str(item)), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1.5)
        pdf.ln(2)

    # Skills
    if skills:
        section_header("Skills")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_x(ml)
        # Check if skills contain categories (colon-separated like "Languages: Python, Java")
        has_categories = any(":" in s for s in skills)
        if has_categories:
            for skill_line in skills:
                clean = ascii_safe(skill_line)
                if ":" in clean:
                    cat, vals = clean.split(":", 1)
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 5, cat.strip() + ":  ", new_x="RIGHT", new_y="TOP")
                    pdf.set_font("Helvetica", "", 9)
                    pdf.multi_cell(content_w - pdf.get_x() + ml, 5, vals.strip(), new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.set_font("Helvetica", "", 9)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 5, clean, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(0.5)
        else:
            skills_text = "  |  ".join(ascii_safe(s) for s in skills)
            pdf.multi_cell(content_w, 5, skills_text, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # Projects
    if projects:
        section_header("Projects")
        for item in projects:
            if isinstance(item, dict):
                proj_title = ascii_safe(item.get("title") or item.get("name", ""))
                tech = ascii_safe(item.get("tech", ""))
                proj_link = ascii_safe(item.get("link", ""))
                if proj_title:
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 5, proj_title, new_x="LMARGIN", new_y="NEXT")
                if tech:
                    pdf.set_font("Helvetica", "I", 8.5)
                    pdf.set_text_color(41, 128, 185)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 4, tech, new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(45, 45, 45)
                if proj_link:
                    pdf.set_font("Helvetica", "", 8)
                    pdf.set_x(ml)
                    pdf.cell(content_w, 4, proj_link, new_x="LMARGIN", new_y="NEXT")
                for bullet in item.get("bullets", []):
                    clean = ascii_safe(re.sub(r'^[\u2022\-*\s]+', '', str(bullet)).strip())
                    if clean:
                        pdf.set_font("Helvetica", "", 9.5)
                        pdf.set_x(ml + 3)
                        pdf.multi_cell(content_w - 5, 4.5, f"- {clean}", new_x="LMARGIN", new_y="NEXT")
                desc = ascii_safe(item.get("description", ""))
                if desc and not item.get("bullets"):
                    pdf.set_font("Helvetica", "", 9)
                    pdf.set_x(ml)
                    pdf.multi_cell(content_w, 4.5, desc, new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.set_x(ml)
                pdf.multi_cell(content_w, 5, ascii_safe(str(item)), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1.5)
        pdf.ln(2)

    # Certifications
    if certifications:
        section_header("Certifications")
        for item in certifications:
            text = ascii_safe(item.get("text", "") if isinstance(item, dict) else str(item))
            if text:
                pdf.set_font("Helvetica", "", 9)
                pdf.set_x(ml)
                pdf.cell(content_w, 5, f"- {text}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    return bytes(pdf.output())


@app.post("/api/set-mode/{session_id}")
async def set_mode(session_id: str, data: dict):
    """Set the session mode to 'resume' or 'ieee'. Runs mode-specific analysis."""
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    mode = data.get("mode", "ieee")
    s["mode"] = mode

    if mode == "resume":
        resume_analysis = analyze_document_for_resume(s["file_text"])
        s["resume_analysis"] = resume_analysis
        s["resume_answers"] = {}
        s["resume_questions_asked"] = []
        s["resume_ready"] = False
        s["resume_text"] = None
        s["resume_html"] = None
        s["resume_data"] = None

        q_result = generate_resume_question(
            s["file_text"], s["resume_answers"], s["resume_questions_asked"], resume_analysis
        )
        if q_result.get("ready"):
            s["resume_ready"] = True
            return {"mode": mode, "ready": True}
        s["_last_resume_qtype"] = q_result.get("type", "")
        return {
            "mode": mode, "ready": False,
            "question": q_result.get("question", ""),
            "options": q_result.get("options", []),
            "context": q_result.get("context", ""),
            "type": q_result.get("type", ""),
        }

    # IEEE mode — run analysis now
    try:
        analysis = analyze_document(s["file_text"])
    except Exception:
        analysis = {}
    s["analysis"] = analysis
    s["answers"] = {}
    s["questions_asked"] = []
    s["ready"] = False

    question = None
    if not analysis.get("ready", False):
        q_result = generate_question(s["file_text"], {}, [], analysis)
        if not q_result.get("ready"):
            question = q_result
    if question:
        s["_last_qtype"] = question.get("type", "")

    return {
        "mode": mode,
        "analysis": analysis,
        "question": question.get("question", "") if question else None,
        "options": question.get("options", []) if question else [],
        "context": question.get("context", "") if question else "",
        "type": question.get("type", "") if question else "",
    }


@app.post("/api/ask-resume/{session_id}")
async def ask_resume_question(session_id: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    q_result = generate_resume_question(
        s["file_text"], s["resume_answers"], s["resume_questions_asked"], s.get("resume_analysis")
    )
    if q_result.get("ready"):
        s["resume_ready"] = True
        return {"ready": True}

    s["_last_resume_qtype"] = q_result.get("type", "")
    return {"question": q_result.get("question", ""), "options": q_result.get("options", []),
            "context": q_result.get("context", ""), "type": q_result.get("type", "")}

@app.post("/api/answer-resume/{session_id}")
async def submit_resume_answer(session_id: str, data: dict):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    question = data.get("question", "")
    answer = data.get("answer", "")
    if not isinstance(question, str):
        return {"error": "question must be a string", "question": "", "options": [], "context": "", "type": ""}
    s["resume_answers"][question] = answer
    s["resume_questions_asked"].append(question)

    qtype = s.get("_last_resume_qtype", "")

    # Handle structured fields
    if qtype in ("name", "name_confirm"):
        if qtype == "name":
            s["resume_analysis"]["name"] = answer.strip()
            s["resume_answers"]["_name_ok"] = True
        elif answer.lower().startswith("y"):
            s["resume_answers"]["_name_ok"] = True
        else:
            # User said no — clear extracted name so the next gate asks for typed input
            s["resume_analysis"]["name"] = ""
            s["resume_answers"]["_name_ok"] = False

    elif qtype in ("email", "email_confirm"):
        if qtype == "email":
            s["resume_analysis"]["email"] = answer.strip()
            s["resume_answers"]["_email_ok"] = True
        elif answer.lower().startswith("y"):
            s["resume_answers"]["_email_ok"] = True
        else:
            s["resume_analysis"]["email"] = ""
            s["resume_answers"]["_email_ok"] = False

    elif qtype == "phone":
        phone = answer.strip()
        if phone and not phone.lower().startswith("n"):
            s["resume_analysis"]["phone"] = phone
        s["resume_answers"]["_phone_ok"] = True

    # Check if user said no more
    user_said_no = any(
        answer.lower().startswith("no") and any(w in question.lower() for w in ("more", "additional", "details"))
        for question in [question]
    )

    try:
        q_result = generate_resume_question(
            s["file_text"], s["resume_answers"], s["resume_questions_asked"], s.get("resume_analysis")
        )
    except Exception:
        q_result = {}
    if q_result.get("ready"):
        s["resume_ready"] = True
        return {"ready": True}

    s["_last_resume_qtype"] = q_result.get("type", "")
    return {"question": q_result.get("question", ""), "options": q_result.get("options", []),
            "context": q_result.get("context", ""), "type": q_result.get("type", "")}


@app.get("/api/generate-resume-stream/{session_id}")
def generate_resume_stream_endpoint(session_id: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    def event_stream():
        resume_text = ""
        try:
            for token in generate_resume_stream(s["file_text"], s["resume_answers"], s.get("resume_analysis", {})):
                resume_text += token
                safe = token.replace("\n", "\\n").replace("\r", "\\r")
                yield f"data: {json.dumps({'type': 'token', 'content': safe})}\n\n"

            s["resume_text"] = resume_text
            resume_data = parse_resume_text(resume_text)
            if s.get("resume_analysis", {}).get("name"):
                resume_data["name"] = s["resume_analysis"]["name"]
            if s.get("resume_analysis", {}).get("email"):
                resume_data["email"] = s["resume_analysis"]["email"]
            if s.get("resume_analysis", {}).get("phone"):
                resume_data["phone"] = s["resume_analysis"]["phone"]

            html = generate_resume_html(resume_data)
            s["resume_html"] = html
            s["resume_data"] = resume_data

            result = {
                "resume_text": resume_text,
                "resume_html": html,
                "resume_data": resume_data,
                "download_html": f"/api/download-resume/{session_id}/html",
                "filename_html": f"{resume_data.get('name', 'Resume').replace(' ', '_')}_Resume.html",
            }
            yield f"data: {json.dumps({'type': 'done', 'result': result})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.post("/api/save-resume/{session_id}")
async def save_resume(session_id: str, data: dict):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    resume_text = data.get("resume_text", s.get("resume_text", ""))
    s["resume_text"] = resume_text
    resume_data = parse_resume_text(resume_text)
    # Overlay contact info from analysis
    if s.get("resume_analysis", {}).get("name"):
        resume_data["name"] = s["resume_analysis"]["name"]
    if s.get("resume_analysis", {}).get("email"):
        resume_data["email"] = s["resume_analysis"]["email"]
    if s.get("resume_analysis", {}).get("phone"):
        resume_data["phone"] = s["resume_analysis"]["phone"]

    html = generate_resume_html(resume_data)
    s["resume_html"] = html
    s["resume_data"] = resume_data

    return {
        "resume_text": resume_text,
        "resume_html": html,
        "resume_data": resume_data,
        "download_html": f"/api/download-resume/{session_id}/html",
        "filename_html": f"{resume_data.get('name', 'Resume').replace(' ', '_')}_Resume.html",
    }


@app.post("/api/edit-resume/{session_id}")
async def edit_resume_endpoint(session_id: str, data: dict):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)
    instruction = (data.get("instruction") or "").strip()
    if not instruction:
        raise HTTPException(400, "No instruction provided")

    current = s.get("resume_text") or ""
    if not current:
        raise HTTPException(400, "No resume to edit yet")

    edited = edit_resume(current, instruction)
    s["resume_text"] = edited
    resume_data = parse_resume_text(edited)
    if s.get("resume_analysis", {}).get("name"):
        resume_data["name"] = s["resume_analysis"]["name"]
    if s.get("resume_analysis", {}).get("email"):
        resume_data["email"] = s["resume_analysis"]["email"]
    if s.get("resume_analysis", {}).get("phone"):
        resume_data["phone"] = s["resume_analysis"]["phone"]

    html = generate_resume_html(resume_data)
    s["resume_html"] = html
    s["resume_data"] = resume_data

    return {
        "resume_text": edited,
        "resume_html": html,
        "resume_data": resume_data,
        "download_html": f"/api/download-resume/{session_id}/html",
        "filename_html": f"{resume_data.get('name', 'Resume').replace(' ', '_')}_Resume.html",
    }


@app.get("/api/preview-resume/{session_id}/{fmt}")
async def preview_resume(session_id: str, fmt: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)
    if fmt == "html" and s.get("resume_html"):
        return Response(content=s["resume_html"], media_type="text/html")
    raise HTTPException(404, "Format not found")

@app.get("/api/download-resume/{session_id}/{fmt}")
async def download_resume(session_id: str, fmt: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    rd = s.get("resume_data") or {}
    name = rd.get("name", "Resume").replace(" ", "_")
    if fmt == "html" and s.get("resume_html"):
        return Response(content=s["resume_html"], media_type="text/html",
                        headers={"Content-Disposition": f"attachment; filename={name}_Resume.html"})
    if fmt == "pdf" and rd:
        pdf_bytes = generate_resume_pdf(rd)
        return Response(content=pdf_bytes, media_type="application/pdf",
                        headers={"Content-Disposition": f"attachment; filename={name}_Resume.pdf"})
    raise HTTPException(404, "Format not found")


@app.post("/api/save-paper/{session_id}")
async def save_paper(session_id: str, data: dict):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    pj = data.get("paper_json", {})
    title = pj.get("title", s["analysis"].get("title", "Research Paper"))
    authors = [
        {"name": a.get("name", ""), "affiliation": a.get("affiliation", ""), "email": a.get("email", "")}
        for a in pj.get("authors", [])
    ]
    abstract = pj.get("abstract", "")
    kw = pj.get("keywords", [])
    if isinstance(kw, str):
        kw = [k.strip() for k in kw.split(",") if k.strip()]
    sections = [{"title": sec.get("title", ""), "content": sec.get("content", "")} for sec in pj.get("sections", [])]
    refs = [{"citation": r} for r in pj.get("references", [])]
    html = generate_ieee_html(title, authors, abstract, sections, kw, s["analysis"].get("domain", "Technology"), refs)
    paper_text = build_markdown(title, abstract, sections, kw, [r["citation"] for r in refs])
    pj["title"] = title
    s["html_content"] = html
    s["paper_text"] = paper_text
    s["paper_json"] = pj
    s["analysis"]["title"] = title
    s["analysis"]["authors"] = authors
    s["analysis"]["keywords"] = kw
    base = title.replace(' ', '_')
    return {
        "paper_text": paper_text,
        "html_content": html,
        "preview_html": f"/api/preview/{session_id}/html",
        "download_html": f"/api/download/{session_id}/html",
        "filename_html": f"{base}.html",
        "download_docx": f"/api/download/{session_id}/docx",
        "filename_docx": f"{base}.docx",
        "paper_json": pj,
    }


@app.post("/api/edit-paper/{session_id}")
async def edit_paper_endpoint(session_id: str, data: dict):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    instruction = (data.get("instruction") or "").strip()
    if not instruction:
        raise HTTPException(400, "No instruction provided")
    current = s.get("paper_text") or ""
    if not current:
        raise HTTPException(400, "No paper to edit yet")
    edited = edit_paper(current, instruction, s["analysis"].get("title", ""))
    s["paper_text"] = edited
    result = parse_paper_text(edited, s["analysis"], session_id)
    s["analysis"]["title"] = result["paper_json"].get("title", s["analysis"].get("title"))
    s["html_content"] = result["html_content"]
    s["paper_json"] = result.get("paper_json")
    return result


@app.get("/api/preview-paper/{session_id}/{fmt}")
async def preview_paper(session_id: str, fmt: str):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)
    if fmt == "html" and s.get("html_content"):
        return Response(content=s["html_content"], media_type="text/html")
    raise HTTPException(404, "Format not found")

@app.get("/api/download-paper/{session_id}/{fmt}")
async def download_paper(session_id: str, fmt: str, format: str = Query("procomm")):
    s = sessions.get(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    _touch_session(session_id)

    paper_fmt = PROCOMM if format != "iemt" else IEMT
    base = s['analysis'].get('title', 'paper').replace(' ', '_')
    if fmt == "html" and s.get("html_content"):
        return Response(content=s["html_content"], media_type="text/html",
                        headers={"Content-Disposition": f"attachment; filename={base}.html"})
    if fmt == "pdf" and s.get("html_content"):
        pdf_bytes = generate_pdf_from_html(s["html_content"], paper_fmt, s.get("images"), s.get("session_dir", ""))
        return Response(content=pdf_bytes, media_type="application/pdf",
                        headers={"Content-Disposition": f"attachment; filename={base}.pdf"})
    if fmt == "docx" and s.get("paper_json"):
        docx_bytes = generate_ieee_docx(s["paper_json"], s.get("session_dir", ""))
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return Response(content=docx_bytes, media_type=media,
                        headers={"Content-Disposition": f'attachment; filename="{base}.docx"'})
    raise HTTPException(404, "Format not found")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
